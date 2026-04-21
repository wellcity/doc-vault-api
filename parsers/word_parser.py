import sys
sys.path.append(".")

import uuid
from docx import Document
from pathlib import Path
from config import IMAGE_DIR
from datetime import datetime, timezone
import json


def parse_docx(file_path: str, file_id: str, metadata: dict | None = None) -> tuple[list[dict], list[str]]:
    """
    解析 Word (.docx) 檔案，回傳 (chunks, image_paths)
    """
    chunks = []
    image_paths = []
    metadata = metadata or {}
    base_name = Path(file_path).stem

    doc = Document(file_path)

    # 建立圖片儲存目錄
    img_dir = Path(IMAGE_DIR) / file_id
    img_dir.mkdir(parents=True, exist_ok=True)

    current_text = ""
    chunk_index = 0
    para_count = 0

    # 遍歷所有段落
    for para in doc.paragraphs:
        para_count += 1
        text = para.text.strip()
        if not text:
            continue

        # 圖片萃取（inline shapes）
        for run in para.runs:
            for shape in run._element.xpath('.//pic:blip', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
                pass  # python-docx 圖片萃取較複雜，先留空

        if len(current_text) + len(text) <= 800:
            current_text += text + "\n"
        else:
            if current_text.strip():
                chunks.append(_make_chunk(file_id, Path(file_path).name, "word", para_count, chunk_index, current_text.strip(), [], metadata))
                chunk_index += 1
            current_text = text + "\n"

    if current_text.strip():
        chunks.append(_make_chunk(file_id, Path(file_path).name, "word", para_count, chunk_index, current_text.strip(), [], metadata))

    # 萃取表格中的文字
    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    table_text += cell_text + " "
        if table_text.strip():
            chunks.append(_make_chunk(file_id, Path(file_path).name, "word", 0, chunk_index, table_text.strip(), [], metadata))
            chunk_index += 1

    return chunks, image_paths


def _make_chunk(file_id: str, source_file: str, file_type: str, page: int, chunk_index: int, text: str, image_paths: list[str], metadata: dict) -> dict:
    return {
        "chunk_id": str(uuid.uuid4()),
        "file_id": file_id,
        "source_file": source_file,
        "file_type": file_type,
        "page": page,
        "chunk_index": chunk_index,
        "text": text,
        "image_paths": image_paths,
        "metadata_json": json.dumps(metadata),
        "embedding": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
