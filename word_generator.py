"""
Word 文件生成器
從大綱結構產生 .docx 檔案（段落、標題、表格、清單）
"""
import sys
sys.path.append(".")

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 配色 ────────────────────────────────────────────────────────────────────

COLOR_PRIMARY = RGBColor(0x1F, 0x49, 0x7D)   # 深藍（標題）
COLOR_ACCENT  = RGBColor(0x2E, 0x86, 0xAB)  # 亮藍（裝飾）
COLOR_TEXT     = RGBColor(0x26, 0x26, 0x26)  # 深灰（內文）
COLOR_MUTED    = RGBColor(0x88, 0x88, 0x88)  # 淺灰（次要）


# ── 工具函式 ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """設定儲存格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """設定儲存格邊框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_run(para, text, bold=False, italic=False, color: RGBColor = None, size: int = None):
    """在段落中加入 Run"""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


# ── 主要生成函式 ────────────────────────────────────────────────────────────

def generate_word_from_outline(outline: dict, output_name: str = "document") -> bytes:
    """
    根據大綱結構產生 Word 文件。

    outline 格式：
    {
        "title": "文件標題",
        "subtitle": "副標題（選填）",
        "author": "作者（選填）",
        "date": "2025/01/01（選填，預設今天）",
        "sections": [
            {
                "heading": "章節標題",
                "content": "內文" or ["項目一", "項目二"],
                "table": [["標題1", "標題2"], ["內容1", "內容2"], ...]（選填）
            },
            ...
        ]
    }
    """
    doc = Document()

    # 頁面設定：左右上下邊界
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    title = outline.get("title", "")
    subtitle = outline.get("subtitle", "")
    author = outline.get("author", "")
    date = outline.get("date", datetime.now().strftime("%Y/%m/%d"))
    sections = outline.get("sections", [])

    # ── 文件標題 ────────────────────────────────────────────────────────────
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = COLOR_PRIMARY
        run.font.name = "Calibri"
        p.space_after = Pt(4)

    # ── 副標題 ──────────────────────────────────────────────────────────────
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT
        run.font.name = "Calibri"
        run.italic = True
        p.space_after = Pt(2)

    # ── 作者 / 日期 ─────────────────────────────────────────────────────────
    if author or date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = " | ".join(filter(None, [author, date]))
        run = p.add_run(meta)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = "Calibri"
        p.space_after = Pt(12)

    # ── 分隔線 ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.space_after = Pt(16)

    # ── 章節內容 ────────────────────────────────────────────────────────────
    for sec in sections:
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        table_data = sec.get("table", [])

        # 章節標題
        if heading:
            p = doc.add_heading(heading, level=1)
            p.runs[0].font.color.rgb = COLOR_PRIMARY
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        # 內文（字串或清單）
        if content:
            if isinstance(content, str):
                p = doc.add_paragraph(content)
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.color.rgb = COLOR_TEXT
                p.runs[0].font.name = "Calibri"
                p.paragraph_format.space_after = Pt(6)
            elif isinstance(content, list):
                for item in content:
                    p = doc.add_paragraph(item, style="List Bullet")
                    if p.runs:
                        p.runs[0].font.size = Pt(11)
                        p.runs[0].font.color.rgb = COLOR_TEXT
                        p.runs[0].font.name = "Calibri"
                    p.paragraph_format.space_after = Pt(3)

        # 表格
        if table_data and len(table_data) >= 2:
            headers = table_data[0]
            rows = table_data[1:]

            tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 標題列
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = str(h)
                _set_cell_bg(hdr_cells[i], "1F497D")
                para = hdr_cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if para.runs:
                    para.runs[0].bold = True
                    para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    para.runs[0].font.size = Pt(10)
                    para.runs[0].font.name = "Calibri"

            # 資料列
            for row_idx, row_data in enumerate(rows):
                row_cells = tbl.rows[row_idx + 1].cells
                bg = "EEF3FA" if row_idx % 2 == 0 else "FFFFFF"
                for col_idx, val in enumerate(row_data):
                    row_cells[col_idx].text = str(val)
                    _set_cell_bg(row_cells[col_idx], bg)
                    para = row_cells[col_idx].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if para.runs:
                        para.runs[0].font.size = Pt(10)
                        para.runs[0].font.color.rgb = COLOR_TEXT
                        para.runs[0].font.name = "Calibri"

            doc.add_paragraph()  # 表格後空一行
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

    # ── 轉成 bytes ──────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
